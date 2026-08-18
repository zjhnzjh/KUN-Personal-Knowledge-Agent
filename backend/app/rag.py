from __future__ import annotations

import hashlib
import json
import math
import re
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Iterable

import httpx

from .config import get_settings
from .database import connect


SUPPORTED = {".pdf", ".docx", ".md", ".xlsx", ".xls", ".png", ".jpg", ".jpeg"}


@dataclass
class ParsedSection:
    locator: str
    heading: str
    text: str


def lexical_text(text: str) -> str:
    """Create explicit Chinese-character and latin-word tokens for SQLite FTS5."""
    return " ".join(re.findall(r"[\u4e00-\u9fff]|[a-zA-Z0-9_]+", text.lower()))


def fts_query(text: str) -> str:
    tokens = lexical_text(text).split()
    return " OR ".join(f'"{token}"' for token in tokens[:64])


def fingerprint(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def parse_document(path: Path, progress: Callable[[int, int], None] | None = None) -> list[ParsedSection]:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED:
        raise ValueError(f"不支持的文件类型：{suffix}")
    if suffix == ".pdf":
        from pypdf import PdfReader
        pages = PdfReader(path).pages
        total = len(pages)
        sections = []
        for index, page in enumerate(pages, 1):
            sections.append(ParsedSection(f"第 {index} 页", "", page.extract_text() or ""))
            if progress and (index == 1 or index == total or index % 5 == 0):
                progress(index, total)
        return sections
    if suffix == ".docx":
        from docx import Document
        doc = Document(path)
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            parts.extend(" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows)
        return [ParsedSection("Word 正文", "", "\n".join(parts))]
    if suffix == ".xlsx":
        from openpyxl import load_workbook
        workbook = load_workbook(path, read_only=True, data_only=False)
        sections: list[ParsedSection] = []
        for sheet in workbook.worksheets:
            lines = []
            for row in sheet.iter_rows(values_only=True):
                values = [str(value).strip() for value in row if value not in (None, "")]
                if values:
                    lines.append(" | ".join(values))
            sections.append(ParsedSection(f"工作表：{sheet.title}", sheet.title, "\n".join(lines)))
        return sections
    if suffix == ".xls":
        import pandas as pd
        sheets = pd.read_excel(path, sheet_name=None, dtype=object)
        return [ParsedSection(
            f"工作表：{name}", name,
            "\n".join(
                " | ".join(str(value).strip() for value in row if str(value).strip() not in {"", "nan"})
                for row in frame.fillna("").itertuples(index=False, name=None)
            ),
        ) for name, frame in sheets.items()]
    if suffix == ".md":
        text = path.read_text(encoding="utf-8", errors="ignore")
        sections, heading, buffer = [], "正文", []
        for line in text.splitlines():
            if line.startswith("#") and buffer:
                sections.append(ParsedSection(heading, heading, "\n".join(buffer)))
                heading, buffer = line.lstrip("# ").strip(), []
            elif line.startswith("#"):
                heading = line.lstrip("# ").strip()
            else:
                buffer.append(line)
        if buffer:
            sections.append(ParsedSection(heading, heading, "\n".join(buffer)))
        return sections
    # Images retain a stable locator. OCR/vision adapters enrich this text asynchronously.
    return [ParsedSection("整张图片", "图片", f"图片文件：{path.name}。等待本地 OCR 与视觉描述。")]


def chunk_sections(sections: Iterable[ParsedSection], size: int = 700, overlap: int = 120) -> list[ParsedSection]:
    chunks: list[ParsedSection] = []
    for section in sections:
        clean = re.sub(r"[ \t]+", " ", section.text).strip()
        if not clean:
            continue
        start = 0
        while start < len(clean):
            end = min(len(clean), start + size)
            if end < len(clean):
                boundary = max(clean.rfind("。", start, end), clean.rfind("\n", start, end))
                if boundary > start + size // 2:
                    end = boundary + 1
            chunks.append(ParsedSection(section.locator, section.heading, clean[start:end]))
            if end >= len(clean):
                break
            start = max(start + 1, end - overlap)
    return chunks


class DashScopeEmbeddings:
    def __init__(self) -> None:
        self.settings = get_settings()

    @property
    def available(self) -> bool:
        return bool(self.settings.dashscope_api_key)

    def encode(self, texts: list[str], text_type: str = "document") -> list[list[float]]:
        if not self.available:
            return []
        last_error: Exception | None = None
        for attempt in range(3):
            try:
                response = httpx.post(
                    f"{self.settings.dashscope_base_url}/embeddings",
                    headers={"Authorization": f"Bearer {self.settings.dashscope_api_key}"},
                    json={"model": self.settings.embedding_model, "input": texts, "dimensions": 1024},
                    timeout=60,
                )
                response.raise_for_status()
                ordered = sorted(response.json()["data"], key=lambda item: item["index"])
                return [item["embedding"] for item in ordered]
            except (httpx.HTTPError, KeyError, ValueError) as error:
                last_error = error
                if attempt < 2:
                    time.sleep(1.5 * (attempt + 1))
        raise RuntimeError("Embedding 服务连续三次请求失败") from last_error


def _cosine(left: list[float], right: list[float]) -> float:
    numerator = sum(a * b for a, b in zip(left, right))
    denominator = math.sqrt(sum(a * a for a in left)) * math.sqrt(sum(b * b for b in right))
    return numerator / denominator if denominator else 0.0


def search(query: str, space_id: str, top_k: int = 5) -> list[dict]:
    query_expression = fts_query(query)
    if not query_expression:
        return []
    with connect() as db:
        lexical = db.execute(
            """SELECT c.id,c.document_id,c.locator,c.heading,c.text,d.title,d.original_name,
                      bm25(chunks_fts) AS score
               FROM chunks_fts JOIN chunks c ON c.id=chunks_fts.chunk_id
               JOIN documents d ON d.id=c.document_id
               WHERE chunks_fts MATCH ? AND d.space_id=? ORDER BY score LIMIT ?""",
            (query_expression, space_id, max(top_k * 3, 12)),
        ).fetchall()
        candidates = {row["id"]: {**dict(row), "lexical_rank": rank} for rank, row in enumerate(lexical, 1)}
        provider = DashScopeEmbeddings()
        if provider.available:
            try:
                query_vectors = provider.encode([query], "query")
            except RuntimeError:
                # The local search path remains usable when the optional provider is
                # unavailable. This is an explicit BM25 fallback, never a fake vector.
                query_vectors = []
            if query_vectors:
                vector_rows = db.execute(
                    """SELECT c.id,c.document_id,c.locator,c.heading,c.text,c.embedding_json,
                              d.title,d.original_name FROM chunks c JOIN documents d ON d.id=c.document_id
                       WHERE d.space_id=? AND c.embedding_json IS NOT NULL""",
                    (space_id,),
                ).fetchall()
                scored = [
                    (row, _cosine(query_vectors[0], json.loads(row["embedding_json"])))
                    for row in vector_rows
                ]
                ranked = sorted(scored, key=lambda item: item[1], reverse=True)[: max(top_k * 3, 12)]
                strongest = ranked[0][1] if ranked else 0.0
                for rank, (row, similarity) in enumerate(ranked, 1):
                    # Dense retrieval always has a "nearest" result. A minimum absolute
                    # similarity and a relative-to-best gate stop unrelated chunks from
                    # being presented as evidence merely because they ranked first.
                    if similarity < 0.35 or (strongest and similarity < strongest * 0.78):
                        continue
                    candidate = candidates.setdefault(row["id"], {**dict(row)})
                    candidate["vector_rank"] = rank
                    candidate["vector_score"] = round(similarity, 5)
        fused = []
        for item in candidates.values():
            rrf = sum(1 / (60 + item[key]) for key in ("lexical_rank", "vector_rank") if key in item)
            fused.append({**item, "score": round(rrf, 5)})
        ranked_fused = sorted(fused, key=lambda item: item["score"], reverse=True)
        results: list[dict] = []
        per_document: dict[str, int] = {}
        for item in ranked_fused:
            document_id = item["document_id"]
            if per_document.get(document_id, 0) >= 2:
                continue
            results.append(item)
            per_document[document_id] = per_document.get(document_id, 0) + 1
            if len(results) >= top_k:
                break
        return results
